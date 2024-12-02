from docx import Document

def create_fir_document(file_name):
  # Create a new Word document
  doc = Document()

  # Add title
  doc.add_heading('Format of the F.I.R.', level=1)
  doc.add_paragraph('\nBook No. ___________\n\nFORM NO. 24.5 (1)\nFIRST INFORMATION REPORT\n')
  doc.add_paragraph('First Information of a Cognizable Crime Reported under Section 154, Criminal Penal Code\n')

  # Add form header
  doc.add_paragraph('Police Station .................................... District .............................. No ..................')
  doc.add_paragraph('Date and hour of Occurrence .....................\n')

  # Add sections
  sections = [
    "1. Date and hour when reported",
    "2. Name and residence of informer and complainant.",
    "3. Brief description of offence (with section) and of property carried off, if any.",
    "4. Place of occurrence and distance and direction from the Police Station.",
    "5. Name & Address of the Criminal.",
    "6. Steps taken regarding investigation explanation of delay in reporting information.",
    "7. Date and Time of dispatch from Police Station."
  ]

  for section in sections:
    doc.add_paragraph(section)

  # Add signature placeholders
  doc.add_paragraph('\nSignature .......................\nDesignation ...............................\n')

  # Add note
  doc.add_paragraph('NOTE: The signature of seal or thumb impression of the informer should be at the end of the information and '
                    'the signature of the Writer of (FIR) should be existed as usual.')

  # Add additional details
  doc.add_paragraph('\nThe above is the format and below is the law on the basis of which the cognizable crime is recorded\n')
  doc.add_paragraph('F.I.R. ON AUTHENTIC INFORMATION\n'
                    'The information given to the Police Officer for registration of a case must be authentic. It should not be gossip but '
                    'should be traced to an individual who should be responsible for imparting information. It may be hearsay but the person '
                    'in possession of hearsay should mention the source of information and take responsibility for it. An irresponsible rumour '
                    'should not result in registration of F.I.R.')

  # Save the document
  doc.save(file_name)
  print(f"{file_name} created successfully!")

# Generate the FIR document
create_fir_document("FIR_Format.docx")